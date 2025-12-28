# brief_interpritator_gui.py
# SAFE brief -> clean Word (DOCX)
# Input:  .docx / .xlsx / .txt (anonymized SAFE brief)
# Output: .docx (clean structured brief)
#
# Dependencies:
#   pip install python-docx openpyxl openai
#
# Works with different OpenAI SDK generations:
# - New: from openai import OpenAI -> client.responses.create(...)
# - Legacy: import openai -> openai.ChatCompletion.create(...)

import os
import re
import json
import datetime as dt
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

try:
    from docx import Document
except Exception:
    Document = None
from openpyxl import load_workbook

try:
    from openai import OpenAI  # new SDK
except Exception:
    OpenAI = None

try:
    import openai as openai_legacy  # legacy SDK
except Exception:
    openai_legacy = None


# -----------------------------
# Helpers: read input
# -----------------------------
def _norm_text(s: str) -> str:
    s = (s or "").replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def read_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_vals = []
            for cell in row.cells:
                c = (cell.text or "").strip()
                if c:
                    row_vals.append(c)
            if row_vals:
                parts.append(" | ".join(row_vals))

    return _norm_text("\n".join(parts))


def read_txt(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            with open(path, "r", encoding=enc) as f:
                return _norm_text(f.read())
        except UnicodeDecodeError:
            continue
    with open(path, "rb") as f:
        return _norm_text(f.read().decode("utf-8", errors="replace"))


def read_xlsx_as_text(path: str) -> str:
    wb = load_workbook(path, data_only=True)
    out = []
    for sh in wb.worksheets:
        out.append(f"[SHEET] {sh.title}")
        max_r = sh.max_row or 0
        max_c = sh.max_column or 0
        for r in range(1, max_r + 1):
            row_vals = []
            empty = True
            for c in range(1, max_c + 1):
                v = sh.cell(r, c).value
                if v is None:
                    row_vals.append("")
                else:
                    empty = False
                    s = str(v).strip()
                    if len(s) > 600:
                        s = s[:600] + "…"
                    row_vals.append(s)
            if not empty:
                while row_vals and row_vals[-1] == "":
                    row_vals.pop()
                if row_vals:
                    out.append(" | ".join(row_vals))
    return _norm_text("\n".join(out))


def try_parse_legitimizer_excel(path: str) -> dict | None:
    try:
        wb = load_workbook(path, data_only=True)
        if "Brief" not in wb.sheetnames:
            return None

        sh = wb["Brief"]
        data = {}
        for r in range(2, (sh.max_row or 0) + 1):
            k = sh.cell(r, 1).value
            v = sh.cell(r, 2).value
            if k is None:
                continue
            key = str(k).strip()
            if not key:
                continue
            val = "" if v is None else str(v).strip()
            data[key] = val

        safe_text = ""
        if "SAFE_TEXT" in wb.sheetnames:
            st = wb["SAFE_TEXT"]
            v = st["A2"].value
            if v:
                safe_text = str(v)

        return {
            "source_type": "legitimizer_excel",
            "brief_table": data,
            "safe_text": _norm_text(safe_text),
        }
    except Exception:
        return None


def build_safe_brief_payload(source_path: str) -> dict:
    ext = os.path.splitext(source_path)[1].lower()

    if ext == ".docx":
        text = read_docx(source_path)
        return {"source_file": os.path.basename(source_path), "safe_raw_text": text, "safe_structured": None}

    if ext == ".txt":
        text = read_txt(source_path)
        return {"source_file": os.path.basename(source_path), "safe_raw_text": text, "safe_structured": None}

    if ext == ".xlsx":
        structured = try_parse_legitimizer_excel(source_path)
        if structured:
            raw = structured["safe_text"] or read_xlsx_as_text(source_path)
            return {
                "source_file": os.path.basename(source_path),
                "safe_raw_text": raw,
                "safe_structured": structured["brief_table"],
            }
        text = read_xlsx_as_text(source_path)
        return {"source_file": os.path.basename(source_path), "safe_raw_text": text, "safe_structured": None}

    raise ValueError("Supported inputs: .docx, .xlsx, .txt")


# -----------------------------
# Prompting + JSON extraction
# -----------------------------
BICM_LAYERS = [
    "BUSINESS_GOAL",
    "RESEARCH_GOAL",
    "DECISION_CONTEXT",
    "SCOPE_AND_BOUNDARIES",
    "TARGET_AUDIENCE",
    "GEOGRAPHY",
    "KEY_QUESTIONS",
    "EXPECTED_OUTPUT",
    "SUCCESS_CRITERIA",
    "CONSTRAINTS",
    "RISKS_AND_ASSUMPTIONS",
]

STATUS_TO_EMOJI = {
    "filled": "🟢 Filled",
    "partial": "🟡 Partial",
    "empty": "🔴 Empty",
}

MATURITY_TO_LABEL = {
    "draft": "🟥 Draft",
    "workable": "🟨 Workable",
    "contract_ready": "🟩 Contract-ready",
}

RU_LAYER_PATHS = {
    "BUSINESS_GOAL": "Цели/Бизнес-цель",
    "RESEARCH_GOAL": "Цели/Исследовательская цель",
    "DECISION_CONTEXT": "Контекст/Решение и зачем",
    "SCOPE_AND_BOUNDARIES": "Объём/Границы и рамки",
    "TARGET_AUDIENCE": "Аудитория/Целевая аудитория",
    "GEOGRAPHY": "География/Регионы",
    "KEY_QUESTIONS": "Вопросы/Ключевые вопросы",
    "EXPECTED_OUTPUT": "Результаты/Ожидаемые материалы",
    "SUCCESS_CRITERIA": "Критерии/Успех",
    "CONSTRAINTS": "Ограничения/Сроки и условия",
    "RISKS_AND_ASSUMPTIONS": "Риски/Допущения",
}


def make_llm_template() -> dict:
    return {
        "layers": [
            {
                "layer": layer,
                "what_is_stated": [{"text": "", "anchor": ""}],
                "gaps": [],
            }
            for layer in BICM_LAYERS
        ]
    }


def make_bicm_template() -> dict:
    return {
        "meta": {"source_file": "", "source_type": "", "generated_at": "", "nda_status": "SAFE"},
        "layers": [
            {
                "layer": layer,
                "what_is_stated": [],
                "gaps": [],
                "status": "empty",
            }
            for layer in BICM_LAYERS
        ],
        "risk_map": {"critical_gaps": [], "interpretation_risks": [], "clear_points": []},
        "maturity": {"level": "draft", "rationale": ""},
    }


def _normalize_anchor(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _ensure_llm_payload(obj: dict) -> dict:
    tpl = make_llm_template()
    out = {"layers": []}

    layers_in = obj.get("layers") if isinstance(obj, dict) else []
    if not isinstance(layers_in, list):
        layers_in = []

    layer_map = {}
    for item in layers_in:
        if not isinstance(item, dict):
            continue
        layer_name = str(item.get("layer", "")).strip().upper()
        if layer_name:
            layer_map[layer_name] = item

    for layer_name in BICM_LAYERS:
        item = layer_map.get(layer_name, {})
        what_is_stated = item.get("what_is_stated", [])
        gaps = item.get("gaps", [])

        if isinstance(what_is_stated, dict):
            what_is_stated = [what_is_stated]
        if isinstance(what_is_stated, str):
            what_is_stated = [{"text": what_is_stated, "anchor": ""}]
        if not isinstance(what_is_stated, list):
            what_is_stated = list(what_is_stated) if what_is_stated else []

        if isinstance(gaps, str):
            gaps = [gaps]
        if not isinstance(gaps, list):
            gaps = list(gaps) if gaps else []

        normalized_items = []
        for entry in what_is_stated:
            if isinstance(entry, str):
                text = entry
                anchor = ""
            elif isinstance(entry, dict):
                text = entry.get("text", "")
                anchor = entry.get("anchor", "")
            else:
                continue
            text = _normalize_anchor(text)
            anchor = _normalize_anchor(anchor)
            if not text and not anchor:
                continue
            normalized_items.append({"text": text, "anchor": anchor})

        gaps = [_normalize_anchor(x) for x in gaps if _normalize_anchor(x)]

        out["layers"].append(
            {
                "layer": layer_name,
                "what_is_stated": normalized_items,
                "gaps": gaps,
            }
        )

    if not out["layers"]:
        out = tpl

    return out


def _derive_layer_status(what_is_stated: list[dict]) -> str:
    count = len([x for x in what_is_stated if str(x.get("text", "") or "").strip()])
    if count >= 3:
        return "filled"
    if count >= 1:
        return "partial"
    return "empty"


def _build_risk_map(layers: list[dict]) -> dict:
    critical = []
    risks = []
    clear = []
    for layer in layers:
        status = layer.get("status")
        name = _ru_layer_path(layer.get("layer", ""))
        if status == "empty":
            critical.append(name)
        elif status == "partial":
            risks.append(name)
        else:
            clear.append(name)
    return {"critical_gaps": critical, "interpretation_risks": risks, "clear_points": clear}


def _derive_maturity(layers: list[dict]) -> dict:
    empty_count = sum(1 for layer in layers if layer.get("status") == "empty")
    partial_count = sum(1 for layer in layers if layer.get("status") == "partial")
    if empty_count > 0:
        return {"level": "draft", "rationale": "Есть критические пробелы."}
    if partial_count > 0:
        return {"level": "workable", "rationale": "Есть интерпретационные риски."}
    return {"level": "contract_ready", "rationale": "Контрактные слои заполнены."}


def build_bicm_report(llm_payload: dict, meta: dict) -> dict:
    tpl = make_bicm_template()
    out = {
        "meta": dict(tpl["meta"]),
        "layers": [],
        "risk_map": dict(tpl["risk_map"]),
        "maturity": dict(tpl["maturity"]),
    }

    out["meta"].update({k: str(meta.get(k, "")).strip() for k in out["meta"].keys()})

    layers_in = llm_payload.get("layers", []) if isinstance(llm_payload, dict) else []
    layer_map = {}
    for item in layers_in:
        if not isinstance(item, dict):
            continue
        layer_name = str(item.get("layer", "")).strip().upper()
        if layer_name:
            layer_map[layer_name] = item

    for layer_name in BICM_LAYERS:
        item = layer_map.get(layer_name, {})
        what_is_stated = item.get("what_is_stated", []) or []
        gaps = item.get("gaps", []) or []
        if not isinstance(what_is_stated, list):
            what_is_stated = []
        if not isinstance(gaps, list):
            gaps = []
        status = _derive_layer_status(what_is_stated)
        out["layers"].append(
            {
                "layer": layer_name,
                "status": status,
                "what_is_stated": what_is_stated,
                "gaps": gaps,
            }
        )

    out["risk_map"] = _build_risk_map(out["layers"])
    out["maturity"] = _derive_maturity(out["layers"])
    return out


def extract_json_strict(text: str) -> dict:
    text = (text or "").strip()
    if not text:
        raise ValueError("Empty model output")

    try:
        return json.loads(text)
    except Exception:
        pass

    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not m:
        raise ValueError("Could not find JSON object in model output")
    return json.loads(m.group(0))


def build_instruction(tone: str) -> str:
    tpl = json.dumps(make_llm_template(), ensure_ascii=False, indent=2)
    layers_text = "\n".join([f"- {layer}" for layer in BICM_LAYERS])
    return (
        "Ты — старший методолог и редактор. Получаешь обезличенный SAFE-бриф.\n"
        "Задача: извлечь ТОЛЬКО смыслы для контрактного отчёта (без оценок).\n"
        "Строгие правила:\n"
        "1) НИЧЕГО не выдумывай. Только то, что явно есть в SAFE-входе.\n"
        "2) Никаких выводов о зрелости/качестве. Никаких оценок.\n"
        "3) Никаких чисел и цифр.\n"
        "4) Никаких брендов, доменных имён, имён людей, точных адресов.\n"
        "   Если встречается — замени токенами BRAND_#, DOMAIN_#, PERSON_#, GEO_#.\n"
        "5) Для каждого слоя: что сказано (text) + короткий anchor (цитата от пяти до двадцати пяти слов).\n"
        "6) Если данных по слою нет — what_is_stated пустой список. gaps может быть пустым.\n"
        "7) Слои строго в таком порядке:\n"
        f"{layers_text}\n"
        "8) Тон: " + (tone or "деловой, чёткий, без лишних слов") + "\n"
        "Верни РОВНО один JSON-объект без пояснений, без обёрток, без markdown.\n"
        "JSON должен иметь такую структуру (поля и типы должны совпадать):\n"
        f"{tpl}\n"
    )


# -----------------------------
# OpenAI call (multi-SDK)
# -----------------------------
def call_llm_to_json(api_key: str, model: str, payload: dict, tone: str) -> dict:
    instruction = build_instruction(tone)

    user_obj = {
        "source_file": payload["source_file"],
        "safe_raw_text": payload["safe_raw_text"],
        "safe_structured_table": payload["safe_structured"] or {},
    }
    user_text = json.dumps(user_obj, ensure_ascii=False)

    # New SDK
    if OpenAI is not None:
        client = OpenAI(api_key=api_key)

        # IMPORTANT FIX: use "input_text" instead of "text"
        try:
            resp = client.responses.create(
                model=model,
                input=[
                    {"role": "system", "content": instruction},
                    {"role": "user", "content": [{"type": "input_text", "text": user_text}]},
                ],
                response_format={"type": "json_object"},
            )
            out_text = getattr(resp, "output_text", None)
            if not out_text:
                out_text = json.dumps(resp.model_dump(), ensure_ascii=False)
            return _ensure_llm_payload(extract_json_strict(out_text))

        except TypeError:
            # response_format not supported
            resp = client.responses.create(
                model=model,
                input=[
                    {"role": "system", "content": instruction},
                    {"role": "user", "content": [{"type": "input_text", "text": user_text}]},
                ],
            )
            out_text = getattr(resp, "output_text", None)
            if not out_text:
                out_text = json.dumps(resp.model_dump(), ensure_ascii=False)
            return _ensure_llm_payload(extract_json_strict(out_text))

        except Exception:
            # fall through to legacy if possible
            pass

    # Legacy SDK
    if openai_legacy is None:
        raise RuntimeError("OpenAI SDK not available. Install: pip install openai")

    openai_legacy.api_key = api_key

    try:
        resp = openai_legacy.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": instruction},
                {"role": "user", "content": user_text},
            ],
            temperature=0.2,
        )
        text = resp["choices"][0]["message"]["content"]
        return _ensure_llm_payload(extract_json_strict(text))
    except AttributeError:
        raise RuntimeError("Your openai package is too old/odd. Upgrade: pip install -U openai")


# -----------------------------
# Word rendering
# -----------------------------
def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]):
    for it in items:
        t = (it or "").strip()
        if not t:
            continue
        doc.add_paragraph(t, style="List Bullet")


def _status_label(status: str) -> str:
    return STATUS_TO_EMOJI.get(status, "🔴 Empty")


def _maturity_label(level: str) -> str:
    return MATURITY_TO_LABEL.get(level, "🟥 Draft")


def _ru_layer_path(layer_name: str) -> str:
    return RU_LAYER_PATHS.get((layer_name or "").strip().upper(), (layer_name or "").strip())


def _layer_comment(layer: dict) -> str:
    status = layer.get("status", "empty")
    if status == "filled":
        return "Заполнено"
    if status == "partial":
        return "Частично заполнено"
    return "Нет данных"


def _format_statement(item: dict) -> str:
    text = _normalize_anchor(item.get("text", ""))
    anchor = _normalize_anchor(item.get("anchor", ""))
    if text and anchor:
        return f"{text} — «{anchor}»"
    if text:
        return text
    if anchor:
        return f"«{anchor}»"
    return ""


def render_docx(out_path: str, doc_json: dict, meta: dict):
    if Document is None:
        raise RuntimeError("нужен python-docx")
    doc = Document()

    add_heading(doc, "BICM Interpretation (SAFE)", level=0)

    add_heading(doc, "1. Паспорт интерпретации", level=1)
    source_file = meta.get("source_file", "")
    source_type = meta.get("source_type", "")
    source_suffix = f" ({source_type})" if source_type else ""
    add_paragraph(doc, f"Источник: {source_file}{source_suffix}")
    add_paragraph(doc, f"Дата генерации: {meta.get('generated_at','')}")
    add_paragraph(doc, "NDA-статус: SAFE")
    add_paragraph(doc, "Цель: приведение к контрактно-читаемому виду для оценки полноты и рисков")

    add_heading(doc, "2. Карта контрактных слоёв", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Слой (RU path)"
    hdr_cells[1].text = "Статус"
    hdr_cells[2].text = "Комментарий"

    for layer in doc_json.get("layers", []):
        row_cells = table.add_row().cells
        row_cells[0].text = _ru_layer_path(layer.get("layer", ""))
        row_cells[1].text = _status_label(layer.get("status", "empty"))
        row_cells[2].text = _layer_comment(layer)

    add_heading(doc, "3. Расшифровка по слоям", level=1)
    for layer in doc_json.get("layers", []):
        layer_name = _ru_layer_path(layer.get("layer", ""))
        add_heading(doc, f"СЛОЙ: {layer_name}", level=2)
        add_paragraph(doc, f"Статус: {_status_label(layer.get('status', 'empty'))}")

        add_paragraph(doc, "Что сказано:")
        what_is_stated = layer.get("what_is_stated", []) or []
        statements = [_format_statement(x) for x in what_is_stated if isinstance(x, dict)]
        statements = [s for s in statements if s]
        if statements:
            add_bullets(doc, statements)
        else:
            add_paragraph(doc, "—")

        add_paragraph(doc, "Пробелы/неясности:")
        gaps = layer.get("gaps", []) or []
        if gaps:
            add_bullets(doc, gaps)
        else:
            add_paragraph(doc, "—")

    add_heading(doc, "4. Risk map", level=1)
    def risk_block(label: str, items: list[str]):
        add_paragraph(doc, label)
        if items:
            add_bullets(doc, items)
        else:
            add_paragraph(doc, "—")

    risk_map = doc_json.get("risk_map", {}) or {}
    risk_block("🔴 Critical gaps", risk_map.get("critical_gaps", []) or [])
    risk_block("🟡 Interpretation risks", risk_map.get("interpretation_risks", []) or [])
    risk_block("🟢 Fixed / clear", risk_map.get("clear_points", []) or [])

    add_heading(doc, "5. Brief maturity", level=1)
    maturity = doc_json.get("maturity", {}) or {}
    rationale = maturity.get("rationale", "") or "—"
    add_paragraph(doc, f"{_maturity_label(maturity.get('level', 'draft'))} {rationale}")

    doc.save(out_path)


# -----------------------------
# GUI
# -----------------------------
class BriefInterpretatorGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("brief_interpritator — SAFE brief → clean Word")
        self.geometry("1180x780")

        self.input_path = tk.StringVar(value="")
        self.api_key = tk.StringVar(value="")
        self.model = tk.StringVar(value="gpt-5.2")
        self.tone = tk.StringVar(value="деловой, чёткий, без лишних слов")

        self.payload = None
        self.generated_json = None

        self._build_ui()

    def _build_ui(self):
        pad = {"padx": 10, "pady": 8}

        box = ttk.LabelFrame(self, text="Input & Model")
        box.pack(fill="x", **pad)

        ttk.Label(box, text="SAFE brief file (.docx / .xlsx / .txt):").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(box, textvariable=self.input_path, width=92, state="readonly").grid(row=0, column=1, sticky="we", padx=8, pady=6)
        ttk.Button(box, text="Choose…", command=self.pick_file).grid(row=0, column=2, padx=8, pady=6)

        ttk.Label(box, text="OpenAI API key:").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(box, textvariable=self.api_key, width=70, show="•").grid(row=1, column=1, sticky="w", padx=8, pady=6)

        ttk.Label(box, text="Model:").grid(row=1, column=2, sticky="e", padx=8, pady=6)
        ttk.Combobox(
            box,
            textvariable=self.model,
            state="readonly",
            width=18,
            values=["gpt-5.2", "gpt-5-mini", "gpt-4.1", "gpt-4o", "gpt-4o-mini"],
        ).grid(row=1, column=3, sticky="w", padx=8, pady=6)

        ttk.Label(box, text="Tone (RU):").grid(row=2, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(box, textvariable=self.tone, width=92).grid(row=2, column=1, columnspan=3, sticky="we", padx=8, pady=6)

        box.columnconfigure(1, weight=1)

        act = ttk.Frame(self)
        act.pack(fill="x", **pad)

        ttk.Button(act, text="Preview input", command=self.preview_input).pack(side="left", padx=8)
        ttk.Button(act, text="Generate BICM (JSON)", command=self.generate).pack(side="left", padx=8)
        ttk.Button(act, text="Export Word (BICM)", command=self.save_word).pack(side="left", padx=8)

        self.status = ttk.Label(act, text="Ready.")
        self.status.pack(side="left", padx=12)

        out = ttk.Notebook(self)
        out.pack(fill="both", expand=True, **pad)

        self.txt_in = tk.Text(out, wrap="word")
        self.txt_json = tk.Text(out, wrap="none")
        self.txt_log = tk.Text(out, wrap="word")

        out.add(self.txt_in, text="Input preview (SAFE)")
        out.add(self.txt_json, text="Generated JSON")
        out.add(self.txt_log, text="Log")

        self._set_text(self.txt_in, "Choose a SAFE brief file to preview.")
        self._set_text(self.txt_json, "")
        self._set_text(self.txt_log, "")

    def _set_text(self, widget: tk.Text, text: str):
        widget.config(state="normal")
        widget.delete("1.0", "end")
        widget.insert("1.0", text)
        widget.config(state="normal")

    def log(self, msg: str):
        self.txt_log.config(state="normal")
        self.txt_log.insert("end", msg + "\n")
        self.txt_log.see("end")
        self.txt_log.config(state="normal")

    def pick_file(self):
        path = filedialog.askopenfilename(
            title="Choose SAFE brief",
            filetypes=[
                ("Word document", "*.docx"),
                ("Excel workbook", "*.xlsx"),
                ("Text file", "*.txt"),
                ("All files", "*.*"),
            ],
        )
        if not path:
            return
        self.input_path.set(path)
        self.status.config(text="File selected.")
        self.payload = None
        self.generated_json = None
        self._set_text(self.txt_json, "")
        self._set_text(self.txt_log, "")
        self.log(f"[OK] Selected: {path}")

    def preview_input(self):
        if not self.input_path.get():
            messagebox.showinfo("No file", "Choose a SAFE brief file first.")
            return
        try:
            self.payload = build_safe_brief_payload(self.input_path.get())
            preview = self.payload["safe_raw_text"]
            if len(preview) > 20000:
                preview = preview[:20000] + "\n\n…(truncated preview)…"
            self._set_text(self.txt_in, preview)
            self.log(f"[OK] Loaded input. chars={len(self.payload['safe_raw_text'])}")
            if self.payload["safe_structured"]:
                self.log("[OK] Detected structured 'Brief' table from legitimizer Excel.")
        except Exception as e:
            messagebox.showerror("Error", str(e))
            self.log(f"[ERROR] preview_input: {e}")

    def generate(self):
        if not self.input_path.get():
            messagebox.showinfo("No file", "Choose a SAFE brief file first.")
            return
        if not self.api_key.get().strip():
            messagebox.showinfo("No API key", "Paste your OpenAI API key.")
            return

        if self.payload is None:
            try:
                self.payload = build_safe_brief_payload(self.input_path.get())
            except Exception as e:
                messagebox.showerror("Error", str(e))
                return

        self.status.config(text="Generating…")
        self.update_idletasks()
        self._set_text(self.txt_log, "")

        try:
            meta = {
                "source_file": os.path.basename(self.input_path.get()),
                "source_type": self._detect_source_type(),
                "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
                "nda_status": "SAFE",
            }
            self.log("[BICM] extracting layers via LLM")
            llm_payload = call_llm_to_json(
                api_key=self.api_key.get().strip(),
                model=self.model.get().strip(),
                payload=self.payload,
                tone=self.tone.get().strip() or "деловой, чёткий, без лишних слов",
            )
            self.log("[BICM] deterministic scoring")
            doc_json = build_bicm_report(llm_payload, meta)
            self.generated_json = doc_json
            self._set_text(self.txt_json, json.dumps(doc_json, ensure_ascii=False, indent=2))
            self.log("[OK] Generated JSON.")
            self.status.config(text="Generated.")
        except Exception as e:
            messagebox.showerror("Generation error", str(e))
            self.log(f"[ERROR] generate: {e}")
            self.status.config(text="Error.")

    def save_word(self, notify: bool = True):
        if not self.generated_json:
            self.log("Nothing to export: generate interpretation first.")
            return

        try:
            if Document is None:
                self.log("Error: нужен python-docx для экспорта Word.")
                return
            out_path = self._ask_export_path()
            if not out_path:
                return
            self.log("[BICM] building Word document")
            render_docx(out_path, self.generated_json, self.generated_json.get("meta", {}))
            self.log(f"[BICM] export completed: {out_path}")
            self.status.config(text="Exported.")
            if notify:
                messagebox.showinfo("Exported", out_path)
        except Exception as e:
            messagebox.showerror("Save error", str(e))
            self.log(f"[ERROR] save_word: {e}")

    def _ask_export_path(self) -> str:
        in_path = self.input_path.get() or "brief"
        base = os.path.splitext(os.path.basename(in_path))[0]
        base = re.sub(r"\s+", " ", base).strip() or "brief"
        initial = f"{base}__BICM.docx"
        return filedialog.asksaveasfilename(
            title="Export BICM Word",
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
            initialfile=initial,
        )

    def _detect_source_type(self) -> str:
        ext = os.path.splitext(self.input_path.get() or "")[1].lower()
        if ext in (".docx", ".xlsx", ".txt"):
            return ext.lstrip(".")
        return ""


def main():
    app = BriefInterpretatorGUI()
    app.mainloop()


if __name__ == "__main__":
    main()
